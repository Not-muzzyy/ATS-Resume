import re
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph):
    """Add a horizontal line under a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def markdown_to_docx(markdown_text: str, candidate_name: str = "") -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = markdown_text.strip().split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            doc.add_paragraph()
            continue

        # H1 — Name
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, bold=True, size=20, color=(0, 70, 127))
            continue

        # H2 — Section headers
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            set_font(run, bold=True, size=12, color=(46, 117, 182))
            add_horizontal_line(p)
            continue

        # H3 — Job title / company / project
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            # Split by | for title | company | date
            parts = [x.strip() for x in text.split('|')]
            if len(parts) >= 2:
                run1 = p.add_run(parts[0])
                set_font(run1, bold=True, size=11)
                run2 = p.add_run('  |  ' + '  |  '.join(parts[1:]))
                set_font(run2, bold=False, size=10, color=(100, 100, 100))
            else:
                run = p.add_run(text)
                set_font(run, bold=True, size=11)
            continue

        # Bullet points
        if line.startswith('- ') or line.startswith('• '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline bold
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.size = Pt(10.5)
            continue

        # Numbered list
        if re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            continue

        # Bold line (standalone)
        if line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run, bold=True, size=11)
            continue

        # Horizontal rule
        if line.startswith('---'):
            continue

        # Regular paragraph with possible inline bold
        p = doc.add_paragraph()
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.size = Pt(10.5)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
